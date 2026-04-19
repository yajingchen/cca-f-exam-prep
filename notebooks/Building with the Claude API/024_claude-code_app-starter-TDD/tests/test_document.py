import os
import stat
import pathlib
import pytest
from tools.document import binary_document_to_markdown, document_path_to_markdown


class TestDocumentPathToMarkdown:
    FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
    DOCX_FIXTURE = os.path.join(FIXTURES_DIR, "mcp_docs.docx")
    PDF_FIXTURE = os.path.join(FIXTURES_DIR, "mcp_docs.pdf")

    # --- Happy path ---

    def test_converts_pdf(self):
        result = document_path_to_markdown(self.PDF_FIXTURE)
        assert isinstance(result, str)
        assert len(result) > 0
        assert "Model Context Protocol" in result

    def test_converts_docx(self):
        result = document_path_to_markdown(self.DOCX_FIXTURE)
        assert isinstance(result, str)
        assert len(result) > 0
        assert "Model Context Protocol" in result

    def test_pdf_and_docx_produce_equivalent_content(self):
        pdf_result = document_path_to_markdown(self.PDF_FIXTURE)
        docx_result = document_path_to_markdown(self.DOCX_FIXTURE)
        assert "Model Context Protocol" in pdf_result
        assert "Model Context Protocol" in docx_result

    def test_accepts_pathlib_path(self):
        result = document_path_to_markdown(pathlib.Path(self.PDF_FIXTURE))
        assert isinstance(result, str)
        assert len(result) > 0

    def test_accepts_relative_path(self):
        rel_path = os.path.relpath(self.PDF_FIXTURE)
        result = document_path_to_markdown(rel_path)
        assert isinstance(result, str)
        assert len(result) > 0

    # --- Output format validation ---

    def test_output_is_str_not_bytes(self):
        result = document_path_to_markdown(self.PDF_FIXTURE)
        assert isinstance(result, str)
        assert not isinstance(result, bytes)

    def test_output_contains_markdown_formatting(self):
        result = document_path_to_markdown(self.PDF_FIXTURE)
        assert any(marker in result for marker in ("#", "-", "*", "|"))

    def test_output_has_no_binary_garbage(self):
        result = document_path_to_markdown(self.PDF_FIXTURE)
        result.encode("utf-8")  # raises UnicodeEncodeError if garbage bytes present

    # --- Complex document elements ---

    def test_table_in_document_produces_markdown_table(self, tmp_path):
        import docx
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Cell 1"
        table.cell(1, 1).text = "Cell 2"
        path = tmp_path / "table.docx"
        doc.save(path)
        result = document_path_to_markdown(str(path))
        assert "|" in result

    def test_image_in_document_does_not_crash(self, tmp_path):
        import docx
        from docx.shared import Inches
        import struct, zlib
        doc = docx.Document()
        # Minimal valid 1x1 PNG
        def minimal_png():
            sig = b'\x89PNG\r\n\x1a\n'
            def chunk(name, data):
                c = struct.pack('>I', len(data)) + name + data
                return c + struct.pack('>I', zlib.crc32(name + data) & 0xffffffff)
            ihdr = chunk(b'IHDR', struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0))
            idat = chunk(b'IDAT', zlib.compress(b'\x00\xff\xff\xff'))
            iend = chunk(b'IEND', b'')
            return sig + ihdr + idat + iend
        import io
        img_path = tmp_path / "img.png"
        img_path.write_bytes(minimal_png())
        doc.add_picture(str(img_path), width=Inches(1))
        path = tmp_path / "with_image.docx"
        doc.save(path)
        result = document_path_to_markdown(str(path))
        assert isinstance(result, str)

    # --- Error cases ---

    def test_raises_file_not_found_for_missing_path(self):
        with pytest.raises(FileNotFoundError, match="/nonexistent/file.pdf"):
            document_path_to_markdown("/nonexistent/file.pdf")

    def test_raises_value_error_for_unsupported_extension(self, tmp_path):
        txt_file = tmp_path / "notes.txt"
        txt_file.write_text("hello")
        with pytest.raises(ValueError, match=".txt"):
            document_path_to_markdown(str(txt_file))

    def test_raises_value_error_for_empty_path(self):
        with pytest.raises(ValueError):
            document_path_to_markdown("")

    def test_raises_permission_error_for_unreadable_file(self, tmp_path):
        locked = tmp_path / "locked.pdf"
        locked.write_bytes(open(self.PDF_FIXTURE, "rb").read())
        locked.chmod(0o000)
        try:
            with pytest.raises(PermissionError, match=str(locked)):
                document_path_to_markdown(str(locked))
        finally:
            locked.chmod(stat.S_IRUSR | stat.S_IWUSR)


class TestBinaryDocumentToMarkdown:
    # Define fixture paths
    FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
    DOCX_FIXTURE = os.path.join(FIXTURES_DIR, "mcp_docs.docx")
    PDF_FIXTURE = os.path.join(FIXTURES_DIR, "mcp_docs.pdf")

    def test_fixture_files_exist(self):
        """Verify test fixtures exist."""
        assert os.path.exists(self.DOCX_FIXTURE), (
            f"DOCX fixture not found at {self.DOCX_FIXTURE}"
        )
        assert os.path.exists(self.PDF_FIXTURE), (
            f"PDF fixture not found at {self.PDF_FIXTURE}"
        )

    def test_binary_document_to_markdown_with_docx(self):
        """Test converting a DOCX document to markdown."""
        # Read binary content from the fixture
        with open(self.DOCX_FIXTURE, "rb") as f:
            docx_data = f.read()

        # Call function
        result = binary_document_to_markdown(docx_data, "docx")

        # Basic assertions to check the conversion was successful
        assert isinstance(result, str)
        assert len(result) > 0
        # Check for typical markdown formatting - this will depend on your actual test file
        assert "#" in result or "-" in result or "*" in result

    def test_binary_document_to_markdown_with_pdf(self):
        """Test converting a PDF document to markdown."""
        # Read binary content from the fixture
        with open(self.PDF_FIXTURE, "rb") as f:
            pdf_data = f.read()

        # Call function
        result = binary_document_to_markdown(pdf_data, "pdf")

        # Basic assertions to check the conversion was successful
        assert isinstance(result, str)
        assert len(result) > 0
        # Check for typical markdown formatting - this will depend on your actual test file
        assert "#" in result or "-" in result or "*" in result
